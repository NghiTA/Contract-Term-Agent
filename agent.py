"""
Contract Term Management Agent
================================
Trợ lý AI giúp Team Legal quản lý vòng đời hợp đồng:
  1. Xác định thời hạn và tính thời điểm hết hiệu lực của hợp đồng.
  2. Cảnh báo trước 15 ngày, phân biệt 2 trường hợp:
        (i)  hết hạn & tự động gia hạn
        (ii) hết hạn & KHÔNG tự động gia hạn
  3. Tiếp tục theo dõi dựa trên các văn bản gia hạn (phụ lục) mới nhất.

RULE CỨNG: Không suy diễn. Chỉ dùng dữ liệu có trong hợp đồng.
Khi không xác định được -> trả về trạng thái KHONG_XAC_DINH ngay.

Triển khai dưới dạng Web API (FastAPI) để deploy lên GreenNode AgentBase.
Gọi model qua GreenNode MaaS (OpenAI-compatible API), mặc định: Minimax.
"""

from __future__ import annotations

import json
import os
import re
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import List, Optional

from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from pydantic import BaseModel, Field

# ---------------------------------------------------------------------------
# Cấu hình (qua biến môi trường – set trên AgentBase Runtime)
# ---------------------------------------------------------------------------
MAAS_BASE_URL = os.getenv("MAAS_BASE_URL", "https://maas.aiplatform.vngcloud.vn/v1")
MAAS_API_KEY = os.getenv("MAAS_API_KEY", "")
MAAS_MODEL = os.getenv("MAAS_MODEL", "minimax")  # Gemma / Qwen / Minimax
CONTRACTS_DIR = os.getenv("CONTRACTS_DIR", "./contracts")
ALERT_WINDOW_DAYS = int(os.getenv("ALERT_WINDOW_DAYS", "15"))
REQUEST_TIMEOUT = int(os.getenv("MAAS_TIMEOUT", "120"))

APP_VERSION = "1.0.0"

# ---------------------------------------------------------------------------
# Đọc văn bản từ nhiều định dạng file (pdf / docx / txt / md)
# ---------------------------------------------------------------------------
def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            import pdfplumber

            parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or "")
            return "\n".join(parts)
        if suffix in (".docx",):
            from docx import Document

            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        if suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        return f"[LỖI ĐỌC FILE {path.name}: {exc}]"
    return ""


SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


def load_contract_documents(contract_path: Path) -> List[dict]:
    """Trả về danh sách document (đã sắp theo thời gian sửa file) cho 1 hợp đồng.

    Quy ước:
      - Mỗi hợp đồng = 1 thư mục con trong CONTRACTS_DIR; mọi file bên trong là
        hợp đồng gốc + các văn bản/phụ lục gia hạn.
      - Hoặc 1 file đơn lẻ = 1 hợp đồng không có văn bản gia hạn.
    """
    docs: List[dict] = []
    if contract_path.is_dir():
        # Sắp theo tên file cho ổn định. KHÔNG dựa vào mtime để suy ra "mới nhất"
        # vì thời gian sửa file không phản ánh đúng văn bản gia hạn mới nhất.
        # Việc xác định văn bản gia hạn mới nhất do model làm dựa trên NGÀY trong nội dung.
        files = sorted(
            [p for p in contract_path.iterdir() if p.suffix.lower() in SUPPORTED_EXTS],
            key=lambda p: p.name,
        )
    else:
        files = [contract_path]
    for f in files:
        docs.append({"filename": f.name, "text": extract_text(f)})
    return docs


def list_contracts(root: Path) -> List[Path]:
    if not root.exists():
        return []
    subdirs = [p for p in root.iterdir() if p.is_dir()]
    if subdirs:
        return sorted(subdirs)
    # fallback: các file phẳng, mỗi file là 1 hợp đồng
    return sorted([p for p in root.iterdir() if p.suffix.lower() in SUPPORTED_EXTS])


# ---------------------------------------------------------------------------
# Gọi LLM qua GreenNode MaaS để trích xuất thông tin có cấu trúc
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """Bạn là trợ lý pháp lý chuyên trích xuất thông tin thời hạn hợp đồng.
NGUYÊN TẮC BẮT BUỘC:
- KHÔNG suy diễn, KHÔNG bịa. Chỉ dùng nội dung có thật trong văn bản.
- Nếu một trường không xác định được rõ ràng, để giá trị null và ghi lý do vào "ghi_chu".
- Khi có nhiều văn bản (hợp đồng gốc + phụ lục/văn bản gia hạn), các văn bản KHÔNG
  được sắp theo thứ tự thời gian. Bạn PHẢI tự xác định văn bản gia hạn MỚI NHẤT dựa
  trên NGÀY ghi TRONG NỘI DUNG từng văn bản (ngày ký, ngày hiệu lực của phụ lục, ngày
  gia hạn đến...), KHÔNG dựa vào thứ tự liệt kê hay tên file.
- Quy trình tính "ngay_het_hieu_luc":
  1) Từ hợp đồng gốc, tính ngày hết hạn ban đầu (ngày hiệu lực + thời hạn).
  2) Nếu có phụ lục/văn bản gia hạn, dùng văn bản gia hạn mới nhất (theo ngày trong
     nội dung) để xác định ngày hết hiệu lực CUỐI CÙNG. Nếu phụ lục ghi "gia hạn đến
     ngày X" thì ngay_het_hieu_luc = X.
  3) Nếu các văn bản mâu thuẫn hoặc không rõ văn bản nào mới nhất → để null và nêu lý
     do trong "ghi_chu".
- Trả về DUY NHẤT một object JSON hợp lệ, không kèm giải thích ngoài JSON.

Định dạng JSON cần trả về:
{
  "ma_hop_dong": string|null,
  "ben_a": string|null,
  "ben_b": string|null,
  "ngay_hieu_luc": "YYYY-MM-DD"|null,
  "thoi_han": string|null,                 // mô tả thời hạn, vd "24 tháng"
  "ngay_het_hieu_luc": "YYYY-MM-DD"|null,   // tính từ văn bản mới nhất
  "tu_dong_gia_han": true|false|null,       // có điều khoản tự động gia hạn?
  "dieu_khoan_gia_han": string|null,        // trích nguyên văn điều khoản gia hạn
  "co_van_ban_gia_han": true|false,         // đã có phụ lục/văn bản gia hạn chưa
  "do_tin_cay": "cao"|"trung_binh"|"thap",
  "ghi_chu": string|null                    // lý do null hoặc điểm cần con người kiểm tra
}"""


def build_user_prompt(contract_name: str, docs: List[dict], today: date) -> str:
    blocks = []
    for i, d in enumerate(docs, 1):
        text = d["text"][:12000]  # giới hạn để tránh vượt context
        blocks.append(f"--- VĂN BẢN {i}: {d['filename']} ---\n{text}")
    joined = "\n\n".join(blocks)
    return (
        f"Hôm nay là {today.isoformat()}.\n"
        f"Hợp đồng: {contract_name}\n"
        f"Số lượng văn bản kèm theo: {len(docs)} (liệt kê theo tên file, "
        f"KHÔNG theo thứ tự thời gian — hãy tự xác định văn bản gia hạn mới nhất "
        f"dựa trên NGÀY ghi trong nội dung).\n\n"
        f"{joined}\n\n"
        "Hãy trích xuất theo đúng định dạng JSON đã quy định."
    )


def call_maas(system: str, user: str) -> str:
    """Gọi GreenNode MaaS qua OpenAI-compatible client."""
    if not MAAS_API_KEY:
        raise RuntimeError(
            "Chưa cấu hình MAAS_API_KEY. Set biến môi trường MAAS_API_KEY trên AgentBase."
        )
    from openai import OpenAI

    client = OpenAI(api_key=MAAS_API_KEY, base_url=MAAS_BASE_URL, timeout=REQUEST_TIMEOUT)
    resp = client.chat.completions.create(
        model=MAAS_MODEL,
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        temperature=0.0,
    )
    return resp.choices[0].message.content or ""


def parse_json_block(raw: str) -> dict:
    """Tách object JSON đầu tiên trong chuỗi trả về của model."""
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()
    match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
    if not match:
        raise ValueError("Model không trả về JSON hợp lệ.")
    return json.loads(match.group(0))


# ---------------------------------------------------------------------------
# Logic nghiệp vụ: tính trạng thái cảnh báo
# ---------------------------------------------------------------------------
def parse_iso_date(value) -> Optional[date]:
    if not value or not isinstance(value, str):
        return None
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(value.strip(), fmt).date()
        except ValueError:
            continue
    return None


def evaluate_alert(extracted: dict, today: date, window: int = ALERT_WINDOW_DAYS) -> dict:
    """Áp dụng rule cảnh báo trên dữ liệu đã trích xuất."""
    expiry = parse_iso_date(extracted.get("ngay_het_hieu_luc"))
    auto_renew = extracted.get("tu_dong_gia_han")

    result = {
        "ngay_het_hieu_luc": expiry.isoformat() if expiry else None,
        "so_ngay_con_lai": None,
        "trang_thai": "KHONG_XAC_DINH",
        "can_canh_bao": False,
        "loai_canh_bao": None,
        "thong_diep": None,
    }

    # RULE CỨNG: không xác định được ngày hết hạn -> phản hồi ngay
    if expiry is None:
        result["thong_diep"] = (
            "KHÔNG xác định được ngày hết hiệu lực từ văn bản. "
            "Cần người phụ trách kiểm tra hợp đồng trực tiếp."
        )
        return result

    days_left = (expiry - today).days
    result["so_ngay_con_lai"] = days_left

    if days_left < 0:
        result["trang_thai"] = "DA_HET_HAN"
    elif days_left <= window:
        result["trang_thai"] = "SAP_HET_HAN"
        result["can_canh_bao"] = True
    else:
        result["trang_thai"] = "CON_HIEU_LUC"

    # Phân loại cảnh báo theo 2 trường hợp tự động gia hạn
    if result["can_canh_bao"] or result["trang_thai"] == "DA_HET_HAN":
        if auto_renew is True:
            result["loai_canh_bao"] = "TU_DONG_GIA_HAN"
            result["thong_diep"] = (
                f"Hợp đồng hết hạn vào {expiry.isoformat()} (còn {days_left} ngày) "
                "và CÓ điều khoản tự động gia hạn. Cần xác nhận với Business xem có "
                "muốn dừng/không gia hạn hay không, và theo dõi thời hạn mới sau gia hạn."
            )
        elif auto_renew is False:
            result["loai_canh_bao"] = "KHONG_TU_DONG_GIA_HAN"
            result["thong_diep"] = (
                f"Hợp đồng hết hạn vào {expiry.isoformat()} (còn {days_left} ngày) "
                "và KHÔNG tự động gia hạn. Cần chủ động ký phụ lục/hợp đồng mới nếu "
                "muốn tiếp tục, tránh để hợp đồng mất hiệu lực."
            )
        else:
            result["loai_canh_bao"] = "GIA_HAN_KHONG_RO"
            result["thong_diep"] = (
                f"Hợp đồng hết hạn vào {expiry.isoformat()} (còn {days_left} ngày) "
                "nhưng KHÔNG xác định rõ có tự động gia hạn hay không. Cần người phụ "
                "trách kiểm tra điều khoản gia hạn."
            )
    else:
        result["thong_diep"] = (
            f"Hợp đồng còn hiệu lực đến {expiry.isoformat()} (còn {days_left} ngày)."
        )
    return result


def analyze_contract(contract_path: Path, today: date) -> dict:
    name = contract_path.name
    docs = load_contract_documents(contract_path)
    if not docs or all(not d["text"].strip() for d in docs):
        return {
            "hop_dong": name,
            "trang_thai": "KHONG_XAC_DINH",
            "loi": "Không đọc được nội dung văn bản nào trong hợp đồng.",
        }
    user_prompt = build_user_prompt(name, docs, today)
    raw = call_maas(SYSTEM_PROMPT, user_prompt)
    try:
        extracted = parse_json_block(raw)
    except Exception as exc:  # noqa: BLE001
        return {
            "hop_dong": name,
            "trang_thai": "KHONG_XAC_DINH",
            "loi": f"Không phân tích được phản hồi của model: {exc}",
            "raw_model_output": raw[:1000],
        }
    alert = evaluate_alert(extracted, today)
    return {
        "hop_dong": name,
        "so_van_ban": len(docs),
        "thong_tin_trich_xuat": extracted,
        "danh_gia_canh_bao": alert,
    }


# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------
app = FastAPI(
    title="Contract Term Management Agent",
    description="Trợ lý AI quản lý vòng đời & cảnh báo hết hạn hợp đồng (GreenNode AgentBase).",
    version=APP_VERSION,
)


class AnalyzeTextRequest(BaseModel):
    contract_name: str = Field(default="Hợp đồng", description="Tên/mã hợp đồng")
    documents: List[str] = Field(
        ..., description="Danh sách nội dung văn bản; phần tử sau là phụ lục/gia hạn mới hơn."
    )
    today: Optional[str] = Field(default=None, description="Ngày tham chiếu YYYY-MM-DD (mặc định hôm nay)")


@app.get("/health")
def health():
    """Health check cho AgentBase Runtime."""
    return {"status": "ok", "version": APP_VERSION, "model": MAAS_MODEL}


@app.get("/scan")
def scan(folder: Optional[str] = None, today: Optional[str] = None):
    """Quét toàn bộ folder hợp đồng và trả về cảnh báo cho từng hợp đồng."""
    root = Path(folder or CONTRACTS_DIR)
    ref_day = parse_iso_date(today) or date.today()
    contracts = list_contracts(root)
    if not contracts:
        return JSONResponse(
            status_code=200,
            content={
                "folder": str(root),
                "ngay_tham_chieu": ref_day.isoformat(),
                "so_hop_dong": 0,
                "ket_qua": [],
                "ghi_chu": f"Không tìm thấy hợp đồng nào trong {root}.",
            },
        )
    results = [analyze_contract(c, ref_day) for c in contracts]
    can_alert = [
        r for r in results
        if r.get("danh_gia_canh_bao", {}).get("can_canh_bao")
        or r.get("danh_gia_canh_bao", {}).get("trang_thai") == "DA_HET_HAN"
    ]
    return {
        "folder": str(root),
        "ngay_tham_chieu": ref_day.isoformat(),
        "so_hop_dong": len(results),
        "so_can_canh_bao": len(can_alert),
        "ket_qua": results,
    }


@app.get("/analyze")
def analyze_one(name: str, folder: Optional[str] = None, today: Optional[str] = None):
    """Phân tích 1 hợp đồng theo tên thư mục/file trong folder."""
    root = Path(folder or CONTRACTS_DIR)
    ref_day = parse_iso_date(today) or date.today()
    target = root / name
    if not target.exists():
        raise HTTPException(status_code=404, detail=f"Không tìm thấy hợp đồng: {name}")
    return analyze_contract(target, ref_day)


@app.post("/analyze-text")
def analyze_text(req: AnalyzeTextRequest):
    """Phân tích hợp đồng từ nội dung text gửi trực tiếp (không cần file)."""
    ref_day = parse_iso_date(req.today) or date.today()
    docs = [{"filename": f"vb_{i+1}", "text": t} for i, t in enumerate(req.documents)]
    if not any(d["text"].strip() for d in docs):
        raise HTTPException(status_code=400, detail="documents rỗng.")
    user_prompt = build_user_prompt(req.contract_name, docs, ref_day)
    raw = call_maas(SYSTEM_PROMPT, user_prompt)
    extracted = parse_json_block(raw)
    alert = evaluate_alert(extracted, ref_day)
    return {
        "hop_dong": req.contract_name,
        "so_van_ban": len(docs),
        "thong_tin_trich_xuat": extracted,
        "danh_gia_canh_bao": alert,
    }


@app.get("/", response_class=HTMLResponse)
def home():
    return f"""<!doctype html>
<html lang="vi"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Contract Term Management Agent</title>
<style>
 body{{font-family:system-ui,Segoe UI,Arial,sans-serif;max-width:820px;margin:40px auto;padding:0 16px;color:#1a2b3c}}
 h1{{color:#0a7d4d}} code{{background:#eef4f0;padding:2px 6px;border-radius:4px}}
 .card{{border:1px solid #d7e3da;border-radius:10px;padding:16px 20px;margin:14px 0}}
 button{{background:#0a7d4d;color:#fff;border:0;padding:10px 16px;border-radius:8px;cursor:pointer;font-size:15px}}
 textarea{{width:100%;min-height:140px;font-family:monospace;padding:8px}}
 pre{{background:#0f1f17;color:#b8f5d4;padding:14px;border-radius:8px;overflow:auto}}
 input{{padding:8px;border:1px solid #ccc;border-radius:6px}}
</style></head><body>
<h1>🛡️ Contract Term Management Agent</h1>
<p>Trợ lý AI quản lý vòng đời hợp đồng — cảnh báo trước <b>{ALERT_WINDOW_DAYS} ngày</b>.
Model: <code>{MAAS_MODEL}</code></p>

<div class="card">
<h3>Quét folder hợp đồng</h3>
<p>Folder mặc định: <code>{CONTRACTS_DIR}</code></p>
<button onclick="scan()">Quét & cảnh báo</button>
</div>

<div class="card">
<h3>Phân tích nhanh từ văn bản</h3>
<p>Mỗi văn bản cách nhau bằng dòng <code>---</code> (văn bản sau là phụ lục mới hơn).</p>
<textarea id="t" placeholder="Dán nội dung hợp đồng vào đây..."></textarea>
<p>Tên hợp đồng: <input id="n" value="Hợp đồng demo"> &nbsp;
Ngày tham chiếu: <input id="d" placeholder="YYYY-MM-DD"></p>
<button onclick="analyze()">Phân tích</button>
</div>

<h3>Kết quả</h3>
<pre id="out">—</pre>

<p style="color:#888">Endpoints: <code>GET /health</code> · <code>GET /scan</code> ·
<code>GET /analyze?name=</code> · <code>POST /analyze-text</code> · <code>/docs</code></p>
<script>
async function scan(){{
  document.getElementById('out').textContent='Đang quét...';
  const r=await fetch('/scan'); document.getElementById('out').textContent=JSON.stringify(await r.json(),null,2);
}}
async function analyze(){{
  const docs=document.getElementById('t').value.split(/^---$/m).map(s=>s.trim()).filter(Boolean);
  const body={{contract_name:document.getElementById('n').value,documents:docs,today:document.getElementById('d').value||null}};
  document.getElementById('out').textContent='Đang phân tích...';
  const r=await fetch('/analyze-text',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify(body)}});
  document.getElementById('out').textContent=JSON.stringify(await r.json(),null,2);
}}
</script>
</body></html>"""


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=int(os.getenv("PORT", "8080")))
